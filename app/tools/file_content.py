"""
Reading and writing file contents.

Separate from filesystem.py, which moves files around without ever looking
inside them. These two tools are what let EDITH summarise a document or save
generated code, rather than only putting it on the clipboard.
"""

from pathlib import Path

from app.security.capabilities import READ, WRITE
from app.tools.filesystem import OpenFileTool, _FileOperationTool

# What a read returns by default. A 16k-character PDF is ~4000 tokens, which
# on its own overflows an 8192-token context once the tool schemas and the
# system prompt are counted. Enough to summarise from; ask for more
# explicitly when a whole file is genuinely needed.
DEFAULT_READ_CHARS = 6_000

# Hard ceiling when max_chars is passed deliberately.
MAX_READ_CHARS = 40_000

# Guard against writing Windows auto-execution artifacts. EDITH is meant to
# produce text and source code; a .bat or .lnk that runs on double-click is a
# different kind of object, and nothing in the spec asks for one.
BLOCKED_WRITE_SUFFIXES = {
    ".exe", ".dll", ".msi", ".com", ".scr", ".sys", ".cpl",
    ".bat", ".cmd", ".ps1", ".psm1", ".vbs", ".vbe", ".wsf", ".wsh",
    ".hta", ".reg", ".lnk", ".pif", ".jar",
}

# Bytes that mean "this is not text". Checked before decoding so a PDF or an
# image is refused rather than returned as mojibake.
_TEXT_HINT_BYTES = 4096


def _extract_pdf(target: Path, limit: int) -> dict:
    """Pull the text layer out of a PDF."""
    try:
        from pypdf import PdfReader

    except ImportError:
        return {
            "success": False,
            "error": (
                "I can't read PDFs: the pypdf package is not installed."
            ),
        }

    try:
        reader = PdfReader(str(target))

        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                return {
                    "success": False,
                    "error": f"{target.name} is password protected.",
                }

        pages = []

        for page in reader.pages:
            pages.append(page.extract_text() or "")

            if sum(len(p) for p in pages) >= limit:
                break

    except Exception as error:
        return {
            "success": False,
            "error": f"I couldn't read {target.name}: {error}",
        }

    text = "\n\n".join(part for part in pages if part.strip())

    if not text.strip():
        # A scanned PDF is images with no text layer. Saying 'it is empty'
        # would be wrong -- there is content, just not text we can extract.
        return {
            "success": False,
            "error": (
                f"{target.name} has no text I can extract -- it looks like "
                f"a scan or images rather than typed text."
            ),
        }

    return {"success": True, "text": text, "pages": len(reader.pages)}


def _extract_docx(target: Path) -> dict:
    try:
        import docx

    except ImportError:
        return {
            "success": False,
            "error": (
                "I can't read Word files: the python-docx package is not "
                "installed."
            ),
        }

    try:
        document = docx.Document(str(target))

    except Exception as error:
        return {
            "success": False,
            "error": f"I couldn't read {target.name}: {error}",
        }

    parts = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]

            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)

    if not text.strip():
        return {"success": False, "error": f"{target.name} has no text in it."}

    return {"success": True, "text": text, "pages": None}


EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
}


def _looks_binary(sample: bytes) -> bool:
    if b"\x00" in sample:
        return True

    if not sample:
        return False

    printable = sum(
        1
        for byte in sample
        if byte in (9, 10, 13) or 32 <= byte < 127 or byte >= 128
    )

    return printable / len(sample) < 0.90


class ReadFileTool(_FileOperationTool):

    name = "read_file"

    description = (
        "Get the contents of a file as text, so you can say what is in it, "
        "summarise it or answer questions about it. Use this whenever the "
        "user asks what a file says or contains -- including PDFs and Word "
        "documents, which this tool extracts text from. Give the full path, "
        "or just part of the file name and optionally the folder."
    )

    parameters = {
        "type": "object",
        "properties": {
            "path": {
                "type": "string",
                "description": (
                    "Full path of the file, or part of its name such as "
                    "'itemized' or 'notes.txt'."
                ),
            },
            "folder": {
                "type": "string",
                "description": (
                    "Optional folder to look in, such as Downloads."
                ),
            },
            "max_chars": {
                "type": "integer",
                "description": (
                    "Maximum characters to return. Defaults to 6000. "
                    "Raise it only when the whole file is needed."
                ),
            },
        },
        "required": ["path"],
        "additionalProperties": False,
    }

    def execute(
        self,
        path: str,
        folder: str = "",
        max_chars: int = DEFAULT_READ_CHARS,
    ):
        found = self._locate(path, folder)

        if isinstance(found, dict):
            return found

        target = found

        limit = max(1, min(int(max_chars or DEFAULT_READ_CHARS), MAX_READ_CHARS))

        # PDFs and Word files are containers, not text. Without an extractor
        # the binary check below refuses them, which reads as 'I can't do
        # that' when the content is perfectly available.
        extractor = EXTRACTORS.get(target.suffix.lower())

        if extractor is not None:
            extracted = (
                extractor(target, limit)
                if target.suffix.lower() == ".pdf"
                else extractor(target)
            )

            if not extracted.get("success"):
                return extracted

            text = extracted["text"]
            truncated = len(text) > limit

            return {
                "success": True,
                "path": str(target),
                "name": target.name,
                "content": text[:limit],
                "lines": text[:limit].count("\n") + 1,
                "characters": len(text[:limit]),
                "pages": extracted.get("pages"),
                "truncated": truncated,
            }

        try:
            with open(target, "rb") as handle:
                sample = handle.read(_TEXT_HINT_BYTES)

                if _looks_binary(sample):
                    return {
                        "success": False,
                        "error": (
                            f"{target.name} is not a text file, so I can't "
                            f"read it."
                        ),
                    }

                # Clamp: a small max_chars can make this negative, and a
                # negative read() means "read everything".
                rest = handle.read(max(0, (limit * 4) - len(sample)))

        except OSError as read_error:
            return {"success": False, "error": str(read_error)}

        text = (sample + rest).decode("utf-8", errors="replace")

        truncated = len(text) > limit

        if truncated:
            text = text[:limit]

        return {
            "success": True,
            "path": str(target),
            "name": target.name,
            "content": text,
            "lines": text.count("\n") + 1,
            "characters": len(text),
            # Say so explicitly: a partial read that reads as complete would
            # let the model answer confidently about what it never saw.
            "truncated": truncated,
        }


    def _locate(self, path: str, folder: str):
        """
        Resolve to a real file. Returns a Path, or an error dict.

        Shares OpenFileTool's search so 'read the itemized pdf' behaves the
        same as 'open the itemized pdf' -- previously read_file had no
        search at all and a bare name resolved to the home folder, which is
        outside every allowed zone.
        """
        path = (path or "").strip().strip('"')

        if not path:
            return {"success": False, "error": "Tell me which file to read."}

        direct, error = self._safe_path(path, mode=READ)

        if error:
            if Path(path).is_absolute() or "/" in path or "\\" in path:
                return error

        else:
            if direct.is_file():
                return direct

            if direct.is_dir():
                return {
                    "success": False,
                    "error": (
                        f"{direct.name} is a folder. Use list_directory to "
                        f"see what is inside it."
                    ),
                }

        matches = OpenFileTool()._search(path, folder)

        if isinstance(matches, dict):
            return matches

        if not matches:
            where = f" in {folder}" if folder else ""

            return {
                "success": False,
                "error": f"I couldn't find a file matching '{path}'{where}.",
            }

        if len(matches) > 1:
            return {
                "success": False,
                "ambiguous": True,
                "query": path,
                "matches": [
                    {"name": item.name, "path": str(item)}
                    for item in matches[:5]
                ],
                "error": f"I found {len(matches)} files matching '{path}'.",
            }

        return matches[0]


class WriteFileTool(_FileOperationTool):

    name = "write_file"

    risk = "destructive"

    description = (
        "Write text to a file, creating it if needed. Use this to save "
        "generated code or notes. Overwriting an existing file asks first. "
        "Set append to true to add to the end instead of replacing."
    )

    parameters = {
        "type": "object",
        "properties": {
            "path": {
                "type": "string",
                "description": (
                    "Full path of the file to write, including the folder."
                ),
            },
            "content": {
                "type": "string",
                "description": (
                    "The exact text to write. Pass the content itself, "
                    "never a description of it."
                ),
            },
            "append": {
                "type": "boolean",
                "description": "Add to the end instead of replacing.",
            },
        },
        "required": ["path", "content"],
        "additionalProperties": False,
    }

    def execute(
        self,
        path: str,
        content: str = "",
        append: bool = False,
        _confirmed: bool = False,
    ):
        if content is None:
            content = ""

        target, error = self._safe_path(path, mode=WRITE, item=True)

        if error:
            return error

        if target.suffix.lower() in BLOCKED_WRITE_SUFFIXES:
            return {
                "success": False,
                "error": (
                    f"I won't create {target.suffix} files -- those run "
                    f"when opened. I can write text or source code instead."
                ),
            }

        if target.is_dir():
            return {
                "success": False,
                "error": f"{target.name} is a folder, not a file.",
            }

        parent = target.parent

        if not parent.is_dir():
            return {
                "success": False,
                "error": (
                    f"The folder does not exist: {parent}. Create it first."
                ),
            }

        # Writing into a folder means the folder must be writable, not just
        # the file path.
        _, error = self._safe_path(str(parent), mode=WRITE)

        if error:
            return error

        existed = target.exists()

        # Replacing a file destroys what was there, and unlike a delete it
        # does not go to the Recycle Bin. Appending adds nothing back.
        if existed and not append and not _confirmed:
            try:
                previous = target.stat().st_size
            except OSError:
                previous = 0

            return {
                "success": False,
                "requires_confirmation": True,
                "summary": (
                    f"{target.name} already exists and holds "
                    f"{previous} bytes. This will replace its contents."
                ),
                "preview": {"count": 1, "names": [target.name]},
            }

        try:
            with open(
                target,
                "a" if append else "w",
                encoding="utf-8",
                newline="",
            ) as handle:
                handle.write(content)

        except OSError as write_error:
            return {"success": False, "error": str(write_error)}

        # Verify rather than assume.
        if not target.exists():
            return {
                "success": False,
                "error": f"The file was not created: {target}",
            }

        return {
            "success": True,
            "path": str(target),
            "name": target.name,
            "characters": len(content),
            "lines": content.count("\n") + 1,
            "appended": bool(append),
            "replaced": existed and not append,
        }
